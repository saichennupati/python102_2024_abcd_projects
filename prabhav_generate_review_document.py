# !pip install python-docx
from docx import Document
import json
import requests

abcd_ids = [711, 733, 523, 542, 731, 763, 669, 739, 690, 742, 724, 112, 668, 714, 671, 32, 722, 434, 713, 751, 760, 102, 654, 695, 762, 470, 33, 432, 429, 265, 673, 111, 588, 734, 692, 477, 518, 720, 445, 437, 752, 206, 670, 544, 491, 737, 685, 492, 53, 676, 575, 469, 183, 717, 665, 715, 410, 425, 430, 586, 769, 235, 462, 741, 524, 718, 758, 465, 745, 520, 431, 721, 171, 484, 759, 672, 433, 749, 502, 756, 747, 757, 26, 765, 667]
abcd_document = Document()

# creating each dictionary and adding it to an empty list
for x in abcd_ids:
    response = requests.get(f'https://abcd2.projectabcd.com/api/getinfo.php?id={x}', headers={"User-Agent": "XY"})
    json_response = json.loads(response.text)
    json_response = dict(json_response)
    if json_response['response_code'] == 400:
        continue
    abcd_document.add_heading(f'{json_response['data']['name']} ABCD ID: {json_response['data']['id']}', level=0)
    abcd_document.add_paragraph(f'Main Description: \n{json_response['data']['description']}')
    abcd_document.add_paragraph(f'Did you know: \n{json_response['data']['did_you_know']}')
    abcd_document.add_heading('Technical Stuff: \n', level=1)
    abcd_document.add_paragraph(f'Category: {json_response['data']['category']}\nType: {json_response['data']['type']}\nState Abbreviation: {json_response['data']['state_name']}\nKey Words: {json_response['data']['key_words']}\nCurrent Status: {json_response['data']['status']}\nNotes: {json_response['data']['notes']}\nBook: {json_response['data']['book']}\nTag Line: {json_response['data']['tag_line']}', style='List Bullet')
    abcd_document.add_page_break()

abcd_document.save('Project_ABCD.docx')
